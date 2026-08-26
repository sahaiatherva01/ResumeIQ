import docx

def generate():
    doc = docx.Document()
    doc.add_heading("Alex Mercer", level=0)
    doc.add_paragraph("Email: alex.mercer@example.com | Phone: +1 555 123 4567 | SF, CA")
    doc.add_paragraph("LinkedIn: linkedin.com/in/alexmercer | GitHub: github.com/alexmercer")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("University of California, Berkeley - B.S. in Computer Science (GPA: 3.82) - 2018 - 2022")

    doc.add_heading("Experience", level=1)
    doc.add_paragraph("TechCorp Solutions - Software Engineer II - 2022 - Present")
    doc.add_paragraph("- Developed and maintained microservices using Python and Flask.")
    doc.add_paragraph("- Reduced page load time by 35% through API query optimizations.")
    doc.add_paragraph("- Spearheaded migration of legacy services, resulting in a 20% cloud hosting cost reduction.")

    doc.add_paragraph("Innovate AI - Software Engineering Intern - 2021")
    doc.add_paragraph("- Built document scanning prototypes achieving 92% classification accuracy.")
    doc.add_paragraph("- Automated reporting pipelines, saving the operations team 5 hours per week.")

    doc.add_heading("Projects", level=1)
    doc.add_paragraph("ResumeLens AI - Tech: Python, Flask, Gemini API, PyMuPDF")
    doc.add_paragraph("- Designed and built a complete single-page interactive report and scoring tool.")

    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("AWS Certified Developer - Associate (2023)")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Flask, Docker, AWS, PostgreSQL, Git, FastAPI, React, JavaScript, HTML, CSS, PyTorch, REST APIs")

    doc.save("test_resume.docx")
    print("test_resume.docx generated successfully!")

if __name__ == "__main__":
    generate()
