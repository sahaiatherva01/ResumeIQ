import io
import re
import fitz  # PyMuPDF
import docx

def detect_sections(text: str) -> dict:
    """
    Detects standard sections in the resume text using common headers.
    Returns a dictionary of section_name -> bool indicating presence.
    """
    sections = {
        "Education": [r"\beducation\b", r"\bacacademic\b", r"\buniversity\b", r"\bschools?\b", r"\bdegrees?\b"],
        "Experience": [r"\bexperience\b", r"\bwork\s+history\b", r"\bemployment\b", r"\bcareer\s+history\b", r"\bprofessional\s+history\b"],
        "Skills": [r"\bskills\b", r"\btechnical\s+skills\b", r"\btechnologies\b", r"\bexpertise\b", r"\bskills?\s+&\s+technologies\b"],
        "Projects": [r"\bprojects\b", r"\bpersonal\s+projects\b", r"\bacademics?\s+projects\b", r"\bkey\s+projects\b"],
        "Certifications": [r"\bcertifications?\b", r"\bcertificates?\b", r"\blicenses?\b", r"\bcredentials?\b"],
        "Achievements": [r"\bachievements?\b", r"\bawards?\b", r"\bhonors?\b", r"\bpublications?\b", r"\baccolades?\b"]
    }
    
    detected = {sec: False for sec in sections.keys()}
    lines = text.split("\n")
    
    for line in lines:
        line_clean = line.strip().lower()
        if len(line_clean) < 50:  # Heading lines are usually short
            for sec, patterns in sections.items():
                if detected[sec]:
                    continue
                for pattern in patterns:
                    if re.search(pattern, line_clean):
                        detected[sec] = True
                        break
                        
    return detected

def parse_pdf(file_bytes: bytes) -> dict:
    """
    Parses a PDF using PyMuPDF (fitz) and analyzes its layout.
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    total_text = []
    has_columns = False
    has_images = False
    
    for page in doc:
        # Extract text blocks
        blocks = page.get_text("blocks")
        # Extract plain text
        page_text = page.get_text()
        total_text.append(page_text)
        
        # Check for images on page
        if len(page.get_images()) > 0:
            has_images = True
            
        # Detect two-column layout
        # If multiple blocks lie parallel horizontally
        rect = page.rect
        width = rect.width
        left_blocks = 0
        right_blocks = 0
        
        for b in blocks:
            x0, y0, x1, y1, text, block_no, block_type = b
            if not text.strip():
                continue
            # Check if block boundaries sit neatly in left or right half
            if x1 < width * 0.53:
                left_blocks += 1
            elif x0 > width * 0.47:
                right_blocks += 1
                
        # Heuristic: if we have several left and right side blocks on a page, it's multi-column
        if left_blocks >= 3 and right_blocks >= 3:
            has_columns = True
            
    full_text = "\n".join(total_text)
    
    # Check if scanned (very little text per page + images)
    avg_chars_per_page = len(full_text.strip()) / max(len(doc), 1)
    is_scanned = (avg_chars_per_page < 150) and has_images
    
    sections = detect_sections(full_text)
    
    return {
        "raw_text": full_text,
        "file_type": "pdf",
        "page_count": len(doc),
        "has_tables": False,  # PyMuPDF doesn't natively check tables simply, we can default to False or let text content capture it
        "has_columns": has_columns,
        "is_scanned": is_scanned,
        "sections_detected": sections,
        "character_count": len(full_text)
    }

def parse_docx(file_bytes: bytes) -> dict:
    """
    Parses a DOCX using python-docx and analyzes its layout.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    text_list = []
    
    # Extract paragraph text
    for p in doc.paragraphs:
        if p.text.strip():
            text_list.append(p.text)
            
    # Extract table text
    has_tables = len(doc.tables) > 0
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            text_list.append(" | ".join(filter(None, row_text)))
            
    full_text = "\n".join(text_list)
    
    # Check columns
    has_columns = False
    for section in doc.sections:
        # If columns are defined on the section and count > 1
        # python-docx section.columns is not always accessible directly if columns aren't configured,
        # but we can try to inspect if cols property exists or count is set.
        try:
            if section.columns and len(section.columns) > 1:
                has_columns = True
        except Exception:
            pass
            
    sections = detect_sections(full_text)
    
    return {
        "raw_text": full_text,
        "file_type": "docx",
        "page_count": 0,  # DOCX doesn't have absolute page count without rendering
        "has_tables": has_tables,
        "has_columns": has_columns or has_tables,  # tables can disrupt parser standard flow
        "is_scanned": len(full_text.strip()) < 50,
        "sections_detected": sections,
        "character_count": len(full_text)
    }

def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Entry point to parse file bytes based on filename extension.
    """
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return parse_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        # Word docs
        return parse_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
